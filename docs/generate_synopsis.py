from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_synopsis():
    doc = Document()

    # Title Section
    title = doc.add_heading('PROJECT SYNOPSIS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run('STREAMHUB: AN ADVANCED MULTI-TIERED SUBSCRIPTION MANAGEMENT SYSTEM WITH REAL-TIME ORCHESTRATION')
    run.bold = True
    run.font.size = Pt(16)

    # Project Abstract
    doc.add_heading('PROJECT ABSTRACT:', level=1)
    doc.add_paragraph(
        "Modern digital services struggle with the complexities of managing user lifecycles, tiered access control, and secure content distribution. "
        "“STREAMHUB” is a high-performance, full-stack web application designed to address these challenges by providing a robust administrative ecosystem "
        "for subscription-based business models. The system focuses on three core pillars: Cinematic User Experience, Production-Grade Security, and Scalable Technical Architecture.\n\n"
        "The application leverages the MERN (MongoDB, Express.js, React, Node.js) stack to deliver a responsive, cloud-ready platform. "
        "Administrators can orchestrate subscription plans (Basic, Premium, Pro) in real-time, while users enjoy a 'Glassmorphism' inspired interface for content discovery. "
        "Data persistence is managed via MongoDB Atlas with Mongoose ODM, ensuring high availability and schema validation. "
        "Security is implemented through a multi-layered approach involving JWT-based Stateless Authentication, Bcrypt password hashing, and Backend Hardening (Rate Limiting, Helmet, and Sanitization)."
    )

    # Project Introduction
    doc.add_heading('PROJECT INTRODUCTION:', level=1)
    doc.add_paragraph(
        "The shift towards 'Software as a Service' (SaaS) and digital streaming has necessitated more sophisticated subscription management tools. "
        "Generic platforms often lack the granular control required for enterprise-level user management and the aesthetic polish expected by modern consumers.\n\n"
        "This project, STREAMHUB, introduces a visionary approach to digital memberships. By utilizing a decoupled architecture (Frontend-as-a-Service and Backend-as-a-Service), "
        "the system ensures maximum modularity and performance. The frontend, powered by React 18 and Vite, offers near-instant loading times and fluid animations using Framer Motion concepts. "
        "The backend follows the MVC (Model-View-Controller) pattern, providing a clean separation of concerns for easier maintenance and future scalability. "
        "STREAMHUB aims to bridge the gap between complex backend logic and high-end visual design, creating a seamless bridge for both end-users and administrators."
    )

    # Scope of the Project
    doc.add_heading('SCOPE OF THE PROJECT:', level=1)
    doc.add_paragraph("The scope encompasses the design and implementation of a secure, production-ready subscription management ecosystem. Key focus areas include:")
    scope_points = [
        "Implementation of Role-Based Access Control (RBAC) specifically for Admin and User tiers.",
        "Development of a secure JWT Vaulting system for cross-session authentication.",
        "Dynamic Plan Orchestration: Allowing admins to create, update, and delete subscription tiers without service downtime.",
        "Immersive Content Delivery: A 'Theater Mode' UI for video playback and interactive media discovery.",
        "Analytical Reporting: Visual dashboards for tracking user registrations, plan popularity, and revenue metrics.",
        "Production Hardening: Mitigating common vulnerabilities like XSS, NoSQL Injection, and Brute Force attacks."
    ]
    for point in scope_points:
        doc.add_paragraph(point, style='List Bullet')

    # Project Modules
    doc.add_heading('LIST OF PROJECT MODULES:', level=1)
    modules = [
        "Authentication & Identity Management (JWT, Bcrypt, Sessioning)",
        "Subscription Engine (Multi-tier logic, Renewal cycles, Plan management)",
        "Discovery Hub (Dynamic gallery, Hover-interactive UI, Metadata mapping)",
        "Media Orchestration (Theater Mode playback, Video state management)",
        "Administrative Command Center (User CRUD, Plan configuration, System logs)",
        "Performance & Analytics Dashboard (Chart.js integration for visual data representation)",
        "Infrastructure & Security Mesh (Node.js Middleware, Rate-limiting, Compression)"
    ]
    for i, module in enumerate(modules, 1):
        doc.add_paragraph(f"{i}. {module}")

    # Hardware/Software requirements
    doc.add_heading('HARDWARE AND SOFTWARE REQUIREMENTS:', level=1)
    
    doc.add_heading('HARDWARE REQUIREMENT:', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Specification'
    
    hardware_reqs = [
        ('Processor', 'Intel Core i5 (11th Gen) or equivalent AMD Ryzen'),
        ('RAM', '8 GB DDR4 or higher'),
        ('Hard Disk', '50 GB SSD (Minimum recommended)'),
        ('Network', 'High-speed broadband for cloud database synchronization')
    ]
    for comp, spec in hardware_reqs:
        row_cells = table.add_row().cells
        row_cells[0].text = comp
        row_cells[1].text = spec

    doc.add_paragraph("") # Space
    
    doc.add_heading('SOFTWARE REQUIREMENT:', level=2)
    table_sw = doc.add_table(rows=1, cols=2)
    table_sw.style = 'Table Grid'
    hdr_cells_sw = table_sw.rows[0].cells
    hdr_cells_sw[0].text = 'Category'
    hdr_cells_sw[1].text = 'Specification'
    
    software_reqs = [
        ('Front-End Environment', 'React.js 18, Vite, Tailwind CSS, Material/Lucide Icons'),
        ('Back-End Infrastructure', 'Node.js LTS, Express.js (RESTful API Design)'),
        ('Database Orchestration', 'MongoDB Atlas (Cloud NoSQL), Mongoose ODM'),
        ('Security Framework', 'JWT, Bcryptjs, Helmet.js, Express-Rate-Limit'),
        ('Development Ecosystem', 'Visual Studio Code, Postman (API Testing), Git/GitHub'),
        ('Operating System', 'Windows 10/11 / Linux (Ubuntu/CentOS recommended for production)')
    ]
    for cat, spec in software_reqs:
        row_cells = table_sw.add_row().cells
        row_cells[0].text = cat
        row_cells[1].text = spec

    # Developer info
    doc.add_heading('PROJECT DEVELOPER DETAILS:', level=1)
    table_dev = doc.add_table(rows=1, cols=3)
    table_dev.style = 'Table Grid'
    hdr_cells_dev = table_dev.rows[0].cells
    hdr_cells_dev[0].text = 'Register Number'
    hdr_cells_dev[1].text = 'Name of the Student'
    hdr_cells_dev[2].text = 'Project Guide'
    
    dev_cells = table_dev.add_row().cells
    dev_cells[0].text = '1DA24MC029'
    dev_cells[1].text = 'MANOJ GOWDA G'
    dev_cells[2].text = 'Dr. Indumathi S. K.'

    doc.add_paragraph("\n\n\nSignature of the Student                                                                                                                                                            ")

    # Save
    doc.save('PROJECT_SYNOPSIS_STREAMHUB_MCA.docx')
    print("Successfully created PROJECT_SYNOPSIS_STREAMHUB_MCA.docx")

if __name__ == "__main__":
    create_synopsis()
